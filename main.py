from PyQt5.QtCore import *
from PyQt5.QtWidgets import *
from PyQt5.QtGui import QIcon
from PyQt5.QtGui import QFont, QPixmap,QPainter, QColor, QBrush
from PyQt5.QtCore import Qt, QTimer
from pgms.indexLayout import Ui_MainWindow
from pgms.database import Database
from pgms.export import CatagoryExporter, AwardedExporter
from docx import Document
import sys, time


class MyApp(QMainWindow, Ui_MainWindow):
    def __init__(self):
        super(MyApp, self).__init__()
        self.setupUi(self)

        self.setWindowIcon(QIcon("assets\\icon.ico"))
        
        # Config
        self.setQTableWidgetConfig()
        self.categoryExportAllBtn.setIcon(QIcon("assets\\word.png"))
        self.categoryExportBtn.setIcon(QIcon("assets\\word.png"))
        self.awardedExportAllBtn.setIcon(QIcon("assets\\word.png"))
        self.awardedExportBtn.setIcon(QIcon("assets\\word.png"))

        # Triggers
        self.fileLoad.triggered.connect(self.loadfile)
        self.deptcomboBox.currentTextChanged.connect(self.categoryComboBoxChanged)
        self.awardedDeptComboBox.currentTextChanged.connect(self.awardedComboBoxChanged)
        self.startYearComboBox.currentTextChanged.connect(self.startYearComboBoxChanged)
        self.endYearComboBox.currentTextChanged.connect(self.endYearComboBoxChanged)
        self.categoryExportAllBtn.clicked.connect(self.categoryExportAll)
        self.categoryExportBtn.clicked.connect(self.categoryExport)
        self.awardedExportBtn.clicked.connect(self.awardedExport)
        self.awardedExportAllBtn.clicked.connect(self.awardedExportAll)
    
    def loadfile(self):
        options = QFileDialog.Options()
        self.filename, _ = QFileDialog.getOpenFileName(self, "Open Excel File", "", "Excel Files (*.xlsx);;All Files (*)", options=options)

        if self.filename.endswith(".xlsx"):
            self.database = Database(self.filename)
            self.sheet_names = self.database.file_content.sheet_names
            self.update_combo_box(self.database.file_content.sheet_names)
        else:
            self.show_warning_message_box("You Have Opened an Unknow file with Unknown sheets")
    
    # Widget Update Functions
    def disconnect_signals(self):
        self.deptcomboBox.currentTextChanged.disconnect(self.categoryComboBoxChanged)
        self.awardedDeptComboBox.currentTextChanged.disconnect(self.awardedComboBoxChanged)
        self.startYearComboBox.currentTextChanged.disconnect(self.startYearComboBoxChanged)
        self.endYearComboBox.currentTextChanged.disconnect(self.endYearComboBoxChanged)

    def reconnect_signals(self):
        self.deptcomboBox.currentTextChanged.connect(self.categoryComboBoxChanged)
        self.awardedDeptComboBox.currentTextChanged.connect(self.awardedComboBoxChanged)
        self.startYearComboBox.currentTextChanged.connect(self.startYearComboBoxChanged)
        self.endYearComboBox.currentTextChanged.connect(self.endYearComboBoxChanged)

    def update_combo_box(self, sheets):
        self.deptcomboBox.clear()
        self.awardedDeptComboBox.clear()
        for sheet in sheets:
            self.deptcomboBox.addItem(sheet)
            self.awardedDeptComboBox.addItem(sheet)
    
    def categoryComboBoxChanged(self, value):
        if value:
            self.database.update_male_female_table(value, self.maleValueTable, self.femaleValueTable)
    
    def awardedComboBoxChanged(self, value):
        if value:
            self.disconnect_signals()
            self.updateYearsComboBox(value)
            self.reconnect_signals()
            start_year = self.startYearComboBox.currentText()
            end_year = self.endYearComboBox.currentText()
            if start_year and end_year:
                start_year = int(start_year)
                end_year = int(end_year)
                self.database.update_awarded_table(value, self.awardedValueTable, start_year, end_year)
            
        
    def startYearComboBoxChanged(self, value):
        if value:
            try:
                sheet_name = self.awardedDeptComboBox.currentText()
                end_year_text = self.endYearComboBox.currentText()
                if end_year_text:
                    end_year = int(end_year_text)
                    self.database.update_awarded_table(sheet_name, self.awardedValueTable, int(value), end_year)
            except ValueError:
                self.show_warning_message_box("Invalid year format. Please select a valid year.")

    def endYearComboBoxChanged(self, value):
        if value:
            try:
                sheet_name = self.awardedDeptComboBox.currentText()
                start_year_text = self.startYearComboBox.currentText()
                if start_year_text:
                    start_year = int(start_year_text)
                    self.database.update_awarded_table(sheet_name, self.awardedValueTable, start_year, int(value))
            except ValueError:
                self.show_warning_message_box("Invalid year format. Please select a valid year.")
        
    def updateYearsComboBox(self, value):
        self.startYearComboBox.clear()
        self.endYearComboBox.clear()
        years_list = self.database.awarded_sort.awardedList.departments[value].years_list
        years_list.sort()
        for year in years_list:
            self.startYearComboBox.addItem(str(year))
            self.endYearComboBox.addItem(str(year))
    
    # Config Functions
    def setQTableWidgetConfig(self):
        self.maleValueTable.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.femaleValueTable.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.awardedValueTable.setEditTriggers(QAbstractItemView.NoEditTriggers)

        self.maleValueTable.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        self.maleValueTable.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        self.awardedValueTable.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        self.awardedValueTable.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)

        self.femaleValueTable.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        self.femaleValueTable.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
    
    def show_warning_message_box(self, msg):
        warning_box = QMessageBox()
        warning_box.setIcon(QMessageBox.Warning)
        warning_box.setWindowTitle("Warning")
        warning_box.setText(msg)
        warning_box.setStandardButtons(QMessageBox.Ok)
        warning_box.exec_()
    
    # Export Functions
    def categoryExportAll(self):
        options = QFileDialog.Options()
        filename, _ = QFileDialog.getSaveFileName(self, "Save Data", "", "Word Files (*.docx);;All Files (*)", options=options)
        if filename:
            document = Document()
            table_file = CatagoryExporter(document)
            
            table_file.add_heading("Male")
            male_table = table_file.create_table()
            for index, sheet in enumerate(self.sheet_names):
                table_file.insert_data(self.database.maleCategoryDataExporter(sheet, index+1), male_table)
            
            table_file.add_heading("\nFemale")
            female_table = table_file.create_table()
            for index, sheet in enumerate(self.sheet_names):
                table_file.insert_data(self.database.femaleCategoryDataExporter(sheet, index+1), female_table)
            document.save(filename)
        
    def categoryExport(self):
        current_sheet_name = self.deptcomboBox.currentText()
        options = QFileDialog.Options()
        filename, _ = QFileDialog.getSaveFileName(self, "Save Data", "", "Word Files (*.docx);;All Files (*)", options=options)
        if filename:
            document = Document()
            table_file = CatagoryExporter(document)

            table_file.add_heading("Male")
            male_table = table_file.create_table()
            table_file.insert_data(self.database.maleCategoryDataExporter(current_sheet_name, 1), male_table)

            table_file.add_heading("\nFemale")
            female_table = table_file.create_table()
            table_file.insert_data(self.database.femaleCategoryDataExporter(current_sheet_name, 1), female_table)
            document.save(filename)
    
    def awardedExport(self):
        options = QFileDialog.Options()
        filename, _ = QFileDialog.getSaveFileName(self, "Save Data", "", "Word Files (*.docx);;All Files (*)", options=options)
        if filename:
            document = Document()
            table_file = AwardedExporter(document)

            current_sheet = self.awardedDeptComboBox.currentText()
            start_year = int(self.startYearComboBox.currentText())
            end_year = int(self.endYearComboBox.currentText())
            years = self.database.awarded_sort.awardedList.departments[current_sheet].years_list

            table_file.add_heading(current_sheet)
            current_table = table_file.create_table()
            
            serial_num = 1

            try:
                for year in range(start_year, end_year+1):
                    if year in years:
                        data = self.database.awardedExporter(current_sheet, year, serial_num)
                        serial_num+=1
                        table_file.insert_data(data, current_table)
            except KeyError:
                exit(1)
            document.save(filename)
    
    def awardedExportAll(self):
        options = QFileDialog.Options()
        filename, _ = QFileDialog.getSaveFileName(self, "Save Data", "", "Word Files (*.docx);;All Files (*)", options=options)
        if filename:
            document = Document()
            table_file = AwardedExporter(document)
            
            
            for sheet in self.sheet_names:
                table_file.add_heading(sheet)
                current_table = table_file.create_table()
                years = self.database.awarded_sort.awardedList.departments[sheet].years_list
                try:
                    for index, year in enumerate(years):
                            data = self.database.awardedExporter(sheet, year, index+1)
                            table_file.insert_data(data, current_table)
                except KeyError:
                    exit(1)
        document.save(filename)

class SplashScreen(QSplashScreen):
    def __init__(self, pixmap):
        super(SplashScreen, self).__init__(pixmap)
        self.setFont(QFont("Arial", 20, QFont.Light))
        self.text = "Loading..."
        self.text_color = Qt.black
        self.background_color = Qt.white  # Change this to the desired background color
        self.alignment = Qt.AlignBottom | Qt.AlignHCenter

    def showMessage(self, text, alignment, color):
        self.text = text
        self.alignment = alignment
        self.text_color = color
        self.repaint()

    def paintEvent(self, event):
        super(SplashScreen, self).paintEvent(event)
        painter = QPainter(self)
        painter.setFont(self.font())
        
        # Measure the text size
        text_rect = painter.boundingRect(self.rect(), self.alignment, self.text)
        
        # Set the background color
        painter.setBrush(QBrush(self.background_color))
        painter.setPen(Qt.NoPen)
        
        # Draw the background rectangle
        painter.drawRect(text_rect)
        
        # Set the text color
        painter.setPen(self.text_color)
        
        # Draw the text
        painter.drawText(self.rect(), self.alignment, self.text)


if __name__ == "__main__":
    app = QApplication(sys.argv)
    img = QPixmap("assets\\splash.png")
    splash = SplashScreen(img)
    screen_rect = app.desktop().screenGeometry()
    splash.move((screen_rect.width() - splash.width()) // 2, (screen_rect.height() - splash.height()) // 2)
    splash.show()
    loading = QTimer()
    loading.singleShot(5000, splash.close)
    time.sleep(3)
    win = MyApp()
    win.show()
    app.exec()