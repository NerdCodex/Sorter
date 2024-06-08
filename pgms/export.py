import pandas as pd
from docx import *
from docx.shared import *
from docx.enum.text import *


class CatagoryExporter:
    def __init__(self, doc):
        self.doc = doc
        self.headings = ["S.NO", "Department", "Religion", "OC", "SC", "ST", "MBC/ONC/BC"]
        self.width_head = [0.54, 1.11, 0.85, 1.28,1.28,1.28,1.28,]
    
    def add_heading(self, Heading):
        heading = self.doc.add_paragraph()
        run = heading.add_run(Heading)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = "Book Antiqua"
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    def create_table(self):
        table = self.doc.add_table(rows=1, cols=7)
        table.style = 'Table Grid'
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                cell.width = Inches(self.width_head[i])
                cell.text = self.headings[i]
                content = cell.paragraphs[0]
                content_run = content.runs[0]
                # ell Content Formatting
                content_run.font.size = Pt(12)
                content_run.font.bold = True
                content.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        return table
    
    def insert_data(self, data, table):
        self.row_1 = table.add_row()
        self.row_2 = table.add_row()
        self.row_3 = table.add_row()

        self.row_1.cells[0].merge(self.row_2.cells[0].merge(self.row_3.cells[0]))
        self.row_1.cells[1].merge(self.row_2.cells[1].merge(self.row_3.cells[1]))
        
        for j in range(7):
            self.row_1.cells[j].text = str(data[0][j])
            if j != 0 and j != 1: 
                self.row_2.cells[j].text = str(data[1][j])
                self.row_3.cells[j].text = str(data[2][j])

class AwardedExporter:
    def __init__(self, doc):
        self.doc = doc
        self.headings = ["S.NO", "Year", "Gender", "Part-Time", "Full-Time", "Total Awarded"]
        self.width_head = [0.54, 1.11, 0.85, 1.23,1.23,1.8]
    
    def add_heading(self, Heading):
        heading = self.doc.add_paragraph()
        run = heading.add_run(Heading)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = "Book Antiqua"
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    def create_table(self):
        table = self.doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                cell.width = Inches(self.width_head[i])
                cell.text = self.headings[i]
                content = cell.paragraphs[0]
                content_run = content.runs[0]
                # ell Content Formatting
                content_run.font.size = Pt(12)
                content_run.font.bold = True
                content.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        return table
    
    def insert_data(self, data, table):
        self.row_1 = table.add_row()
        self.row_2 = table.add_row()
        

        self.row_1.cells[0].merge(self.row_2.cells[0])
        self.row_1.cells[1].merge(self.row_2.cells[1])
        self.row_1.cells[5].merge(self.row_2.cells[5])
        
        for j in range(6):
            self.row_1.cells[j].text = str(data[0][j])
            if j != 0 and j != 1 and j != 5: 
                self.row_2.cells[j].text = str(data[1][j])

if __name__ == "__main__":
    d = Document()
    e = AwardedExporter(d)
    e.add_heading("Tamil\n")
    t = e.create_table()
    e.insert_data([[1, 2020, "Male", 1, 1, 4], [2, 2020, "Female", 1, 1, 4]],t)
    d.save("word.docx")